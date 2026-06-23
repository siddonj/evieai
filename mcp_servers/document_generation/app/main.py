"""Document Generation MCP Server — Multifamily & Brokerage report templates + export (Excel/Word/PDF)."""
from __future__ import annotations

import io
import re
from typing import Any

from fastapi import FastAPI, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel

# Eager imports so Azure Container Apps doesn't time out on first request while
# packages are being loaded. weasyprint in particular scans fonts on first import
# which can take 10-30s on a cold container.
try:
    import docx as _docx  # noqa: F401
    import openpyxl as _openpyxl  # noqa: F401
    import weasyprint as _weasyprint  # noqa: F401
except ImportError:
    pass  # missing libs surface as HTTPException 500 when the endpoint is called

app = FastAPI(title="mcp-document-generation", version="0.3.0")

# ═══════════════════════════════════════════════════════════════════════
#  DEMO DATA  —  Multifamily & Brokerage Document Templates
# ═══════════════════════════════════════════════════════════════════════

_DOCUMENT_TEMPLATES: list[dict[str, Any]] = [
    {
        "id": "doc-mf-001",
        "type": "portfolio_summary",
        "title": "Q2 2026 Portfolio Performance Summary",
        "generated_at": "2026-05-19T10:00:00Z",
        "author": "ResiQ Brokerage Analytics",
        "status": "Generated",
        "pages": 4,
        "word_count": 1450,
        "sections": [
            {
                "heading": "Portfolio Overview",
                "content": "The Memphis multifamily portfolio consists of 8 properties totaling 1,248 units with 92.8% occupancy. Total portfolio value estimated at $178.9M with aggregate NOI of $11.0M. Average cap rate of 6.0% across the portfolio. Rent growth of 4.8% YoY outpaces the Memphis MSA average of 3.2%.",
                "key_metrics": [
                    {"label": "Properties", "value": "8", "trend": "1,248 units"},
                    {"label": "Portfolio Value", "value": "$178.9M", "trend": "+12% YoY"},
                    {"label": "Occupancy", "value": "92.8%", "trend": "+1.2pp QoQ"},
                    {"label": "Avg Rent", "value": "$1,333", "trend": "+4.8% YoY"},
                    {"label": "Total NOI", "value": "$11.0M", "trend": "+8.2% TTM"},
                    {"label": "Avg Cap Rate", "value": "6.0%", "trend": "Stable"},
                ],
            },
            {
                "heading": "Class A Assets — Downtown & Germantown",
                "content": "The Emerson (312 units, $1,895 avg rent, 5.2% cap) and The Vue at Madison (240 units, $1,625 avg rent, 5.8% cap) represent the top of the portfolio. Combined value: $98.6M. These institutional-quality assets are attracting Blackstone and Invesco-level buyers. Average occupancy of 94.1% with 6.2% YoY rent growth in the downtown submarket.",
                "key_metrics": [
                    {"label": "The Emerson", "value": "$66.3M", "trend": "312 units, 5.2% cap"},
                    {"label": "The Vue at Madison", "value": "$32.3M", "trend": "240 units, 5.8% cap"},
                    {"label": "Village at Germantown", "value": "$23.1M", "trend": "156 units, Under Contract"},
                ],
            },
            {
                "heading": "Value-Add & Workforce Housing",
                "content": "Oakwood Crossings (96 units, $1,095 rent, 6.2% cap) presents the clearest value-add opportunity with current rents 18% below market. Highland Ridge (64 units, $975 rent, 6.8% cap) benefits from University of Memphis proximity. Southgate Village (128 units, $845 rent, 7.5% cap) maintains 98% occupancy with stable Section 8 voucher income.",
                "key_metrics": [
                    {"label": "Oakwood Crossings", "value": "$10.5M", "trend": "6.2% cap, value-add"},
                    {"label": "Highland Ridge", "value": "$5.6M", "trend": "6.8% cap, near U of M"},
                    {"label": "Southgate Village", "value": "$6.4M", "trend": "7.5% cap, workforce"},
                ],
            },
            {
                "heading": "Market Context",
                "content": "Memphis MSA continues to benefit from supply constraints — new deliveries down 15% YoY to 420 units. Job growth at 2.1% exceeds national average. Class A cap rates compressing to 5.0-5.8% as institutional capital targets Sunbelt secondary markets. Downtown submarket strongest with 6.2% YoY rent growth.",
                "key_metrics": [
                    {"label": "Market Occupancy", "value": "91.5%", "trend": "+0.8pp QoQ"},
                    {"label": "New Supply", "value": "420 units", "trend": "-15% YoY"},
                    {"label": "Class A Cap Rate", "value": "5.0-5.8%", "trend": "Tightening"},
                    {"label": "Job Growth", "value": "+2.1%", "trend": "Above national"},
                ],
            },
        ],
        "action_items": [
            "Schedule seller update for Village at Germantown — closing Q3 2026",
            "Prepare Oakwood Crossings value-add business plan for buyer tour follow-up",
            "Monitor Invesco LOI progress on The Emerson — backup offers if needed",
            "Update portfolio OM with Q2 2026 financial data by June 1",
        ],
        "tags": ["portfolio", "quarterly", "performance", "memphis"],
    },
    {
        "id": "doc-mf-002",
        "type": "offering_memorandum",
        "title": "Offering Memorandum — The Vue at Madison",
        "generated_at": "2026-05-19T14:00:00Z",
        "author": "ResiQ Brokerage",
        "status": "Draft",
        "pages": 18,
        "word_count": 3200,
        "sections": [
            {
                "heading": "Executive Summary",
                "content": "The Vue at Madison is a 240-unit Class A multifamily asset in downtown Memphis. Built in 2018, the property offers studio, one, and two-bedroom units with an average rent of $1,625. Current NOI of $1.87M at a 5.8% cap rate. Listed at $35.8M ($149K/unit). Current occupancy: 95% (228 of 240 units). Blackstone Real Estate under contract at $34M in due diligence.",
                "key_metrics": [
                    {"label": "Total Units", "value": "240", "trend": "Studio, 1BR, 2BR mix"},
                    {"label": "Year Built", "value": "2018", "trend": "Class A construction"},
                    {"label": "NOI", "value": "$1.87M", "trend": "5.8% cap rate"},
                    {"label": "List Price", "value": "$35.8M", "trend": "$149K/unit"},
                    {"label": "Occupancy", "value": "95%", "trend": "228 units occupied"},
                ],
            },
            {
                "heading": "Location & Demographics",
                "content": "Located at 123 Madison Ave in downtown Memphis's thriving core. Walking distance to Beale Street, FedExForum, AutoZone Park, and the Mississippi Riverfront. 5-minute drive to I-40/I-55. 15-minute drive to Memphis International Airport. Average household income within 3-mile radius: $82K. Population growth: +8% since 2020. Walk score: 92.",
                "key_metrics": [
                    {"label": "Walk Score", "value": "92", "trend": "Walker's Paradise"},
                    {"label": "3-Mile HH Income", "value": "$82K", "trend": "+12% since 2020"},
                    {"label": "Population Growth", "value": "+8%", "trend": "Since 2020"},
                    {"label": "Avg Rent Premium", "value": "22%", "trend": "vs market avg"},
                ],
            },
            {
                "heading": "Financial Summary",
                "content": "Stabilized NOI of $1.87M with 15% upside through rent growth and expense optimization. Current average rent $1,625 vs market comps of $1,850 ($225/unit upside). Expense ratio of 38% vs market average of 42%. Capital reserve funded at $250/unit/year. Roof (2018), HVAC (2018), parking garage (2018) — all near-new with minimal deferred maintenance.",
                "key_metrics": [
                    {"label": "Current NOI", "value": "$1.87M", "trend": "5.8% cap"},
                    {"label": "Pro Forma NOI", "value": "$2.15M", "trend": "6.7% cap stabilized"},
                    {"label": "Rent Upside", "value": "$225/unit", "trend": "14% below market"},
                    {"label": "Expense Ratio", "value": "38%", "trend": "vs 42% market"},
                ],
            },
            {
                "heading": "Unit Mix & Amenities",
                "content": "Mix: 48 studios ($1,295), 120 one-bedrooms ($1,595), 72 two-bedrooms ($1,895). Unit features: quartz countertops, stainless steel appliances, in-unit washer/dryer, smart locks, walk-in closets. Amenities: resort-style pool, fitness center, rooftop lounge with skyline views, covered parking, pet park with dog spa, co-working lounge, 24/7 package lockers.",
                "key_metrics": [
                    {"label": "Studio", "value": "48 units", "trend": "$1,295 avg"},
                    {"label": "1BR", "value": "120 units", "trend": "$1,595 avg"},
                    {"label": "2BR", "value": "72 units", "trend": "$1,895 avg"},
                    {"label": "Pet Units", "value": "60%", "trend": "Pet park + dog spa"},
                ],
            },
        ],
        "action_items": [
            "Update OM with final due diligence findings from purchaser's Phase I",
            "Prepare rent comp exhibits for Blackstone underwriting team",
            "Confirm closing timeline with title company for August 31 target",
        ],
        "tags": ["offering memorandum", "the vue", "downtown", "class a"],
    },
    {
        "id": "doc-mf-003",
        "type": "market_survey",
        "title": "Memphis Multifamily Market Survey — Q2 2026",
        "generated_at": "2026-05-19T09:00:00Z",
        "author": "ResiQ Research",
        "status": "Generated",
        "pages": 6,
        "word_count": 2100,
        "sections": [
            {
                "heading": "Market Overview",
                "content": "Memphis multifamily market continues to strengthen with 91.5% average occupancy (+0.8pp QoQ) and 3.2% YoY rent growth. Supply constraints support landlord fundamentals — only 420 new units delivered in 2025 vs 720 peak in 2022. Job growth of 2.1% driven by logistics (FedEx, Amazon), healthcare (St. Jude, Methodist), and technology sectors.",
                "key_metrics": [
                    {"label": "Market Occupancy", "value": "91.5%", "trend": "+0.8pp QoQ"},
                    {"label": "Avg Rent", "value": "$1,285", "trend": "+3.2% YoY"},
                    {"label": "New Supply", "value": "420 units", "trend": "-15% YoY"},
                    {"label": "Absorption", "value": "380 units", "trend": "92% leased"},
                ],
            },
            {
                "heading": "Submarket Analysis — Downtown",
                "content": "Downtown Memphis (38103) is the strongest submarket with 93.8% occupancy and 6.2% YoY rent growth. The Emerson (opened 2021) and The Vue (2018) anchor the luxury segment with rents averaging $1,750+. New supply limited to adaptive reuse projects. Walkable amenities, riverfront access, and entertainment district drive demand from young professionals.",
                "key_metrics": [
                    {"label": "Downtown Occupancy", "value": "93.8%", "trend": "+1.5pp QoQ"},
                    {"label": "Downtown Avg Rent", "value": "$1,750+", "trend": "+6.2% YoY"},
                    {"label": "Luxury Rent Premium", "value": "+42%", "trend": "vs portfolio avg"},
                    {"label": "Downtown Population", "value": "12,400", "trend": "+18% since 2020"},
                ],
            },
            {
                "heading": "Submarket Analysis — East Memphis / Germantown",
                "content": "East Memphis (38117, 38119) and Germantown (38138) represent stable suburban submarkets with 92.1% occupancy and 2.8% YoY rent growth. Strong school districts and corporate employment anchors. Average rent of $1,375 for Class B and $1,650 for Class A. Value-add opportunities in aging 1970s-1990s garden-style product.",
                "key_metrics": [
                    {"label": "Suburban Occupancy", "value": "92.1%", "trend": "+0.5pp QoQ"},
                    {"label": "Suburban Avg Rent", "value": "$1,375", "trend": "+2.8% YoY"},
                    {"label": "Value-Add Spread", "value": "18%", "trend": "Below market avg"},
                    {"label": "Avg Days on Market", "value": "24", "trend": "-3 days YoY"},
                ],
            },
            {
                "heading": "Sales & Cap Rate Trends",
                "content": "Class A cap rates compressed to 5.0-5.8% as institutional capital targets Memphis for higher yields vs primary markets (NYC 4.0%, LA 3.8%). Class B/C at 6.5-7.5% stable. Dollar volume Q1 2026: $185M (+35% YoY). Average price/unit: $145K Class A, $78K Class B, $52K Class C. Buyer mix: 40% institutional, 35% private equity, 25% local/regional.",
                "key_metrics": [
                    {"label": "Class A Cap Rate", "value": "5.0-5.8%", "trend": "Tightening"},
                    {"label": "Class B/C Cap Rate", "value": "6.5-7.5%", "trend": "Stable"},
                    {"label": "Q1 Sales Volume", "value": "$185M", "trend": "+35% YoY"},
                    {"label": "Avg Price/Unit A", "value": "$145K", "trend": "Class A"},
                ],
            },
        ],
        "action_items": [
            "Update market comps for The Emerson offering memorandum",
            "Prepare East Memphis value-add analysis for Oakwood Crossings buyer",
            "Monitor downtown luxury supply pipeline for competitive positioning",
        ],
        "tags": ["market survey", "memphis", "quarterly", "submarket"],
    },
    {
        "id": "doc-mf-004",
        "type": "broker_price_opinion",
        "title": "Broker Price Opinion — Oakwood Crossings",
        "generated_at": "2026-05-19T11:30:00Z",
        "author": "Linda Thornton, Cushman & Wakefield",
        "status": "Draft",
        "pages": 8,
        "word_count": 1800,
        "sections": [
            {
                "heading": "Property Overview",
                "content": "Oakwood Crossings is a 96-unit garden-style multifamily property at 4525 Oakwood Dr, Memphis, TN 38117. Built 1985 on 5.1 acres. Unit mix: 1BR/1BA (32 units, 685 SF), 2BR/1BA (40 units, 895 SF), 2BR/2BA (24 units, 1,050 SF). Currently 82 units occupied (85.4% occupancy). Average rent $1,095. Condition: B- (40% renovated, deferred maintenance on remaining 60%).",
                "key_metrics": [
                    {"label": "Total Units", "value": "96", "trend": "85.4% occupied"},
                    {"label": "Avg Rent", "value": "$1,095", "trend": "18% below market"},
                    {"label": "Year Built", "value": "1985", "trend": "Garden-style"},
                    {"label": "Lot", "value": "5.1 acres", "trend": "5.3 units/acre"},
                ],
            },
            {
                "heading": "Comparable Sales Analysis",
                "content": "Three recent comparable sales: (1) Highland Ridge (64 units, $5.6M, $87.5K/unit, 6.8% cap) — similar vintage and condition. (2) Poplar Pointe (72 units, $9.5M, $131.9K/unit, 5.9% cap) — renovated, better East Memphis location. (3) Midsouth Village (112 units, $8.2M, $73.2K/unit, 7.2% cap) — unrenovated, lower quality. Oakwood Crossings positioned between Highland Ridge and Poplar Pointe at an estimated $10.5-12.5M.",
                "key_metrics": [
                    {"label": "Highland Ridge", "value": "$87.5K/unit", "trend": "6.8% cap"},
                    {"label": "Poplar Pointe", "value": "$131.9K/unit", "trend": "5.9% cap, renovated"},
                    {"label": "Midsouth Village", "value": "$73.2K/unit", "trend": "7.2% cap, unrenovated"},
                    {"label": "Estimated Value", "value": "$10.5-12.5M", "trend": "$109-130K/unit"},
                ],
            },
            {
                "heading": "Value-Add Analysis",
                "content": "Current NOI of $648K at 6.2% pro forma cap rate. After $1.2M in capital improvements (roof, HVAC replacements, interior renovations in 60% unrenovated units), pro forma NOI of $845K achievable. Stabilized value: $12.8M at 6.6% cap. Value creation: $2.3M above current value. Renovation scope: new countertops, appliances, LVP flooring, lighting, fixtures ($18.7K/unit average).",
                "key_metrics": [
                    {"label": "Current NOI", "value": "$648K", "trend": "6.2% pro forma cap"},
                    {"label": "Pro Forma NOI", "value": "$845K", "trend": "+$197K upside"},
                    {"label": "Renovation Cost", "value": "$1.2M", "trend": "$18.7K/unit"},
                    {"label": "Stabilized Value", "value": "$12.8M", "trend": "$2.3M value-add"},
                ],
            },
            {
                "heading": "Recommendation",
                "content": "Recommended list price: $12.5M ($130K/unit) with a sell range of $11.0-12.5M. Optimal buyer profile: value-add private equity fund or experienced local operator. Target cap rate: 6.0-6.5%. Marketing strategy: targeted outreach to 1031 exchange buyers and SE value-add funds. Estimated DOM: 45-60 days at $12.5M list. Commission: 3% ($375K at $12.5M).",
                "key_metrics": [
                    {"label": "Recommended List", "value": "$12.5M", "trend": "$130K/unit"},
                    {"label": "Sell Range", "value": "$11.0-12.5M", "trend": "6.0-6.5% cap"},
                    {"label": "Est. DOM", "value": "45-60 days", "trend": "At $12.5M list"},
                    {"label": "Commission", "value": "$375K", "trend": "3% at list price"},
                ],
            },
        ],
        "action_items": [
            "Prepare full OM with renovation pro forma for buyer distribution",
            "Identify 1031 exchange buyers with Q3 2026 deadline",
            "Engage contractor for renovation cost verification",
            "Schedule BPO presentation with Oakwood Holdings LLC",
        ],
        "tags": ["broker price opinion", "bpo", "oakwood", "value-add"],
    },
    {
        "id": "doc-mf-005",
        "type": "commission_report",
        "title": "Q2 2026 Commission & Activity Report",
        "generated_at": "2026-05-19T16:00:00Z",
        "author": "ResiQ Brokerage Systems",
        "status": "Generated",
        "pages": 2,
        "word_count": 850,
        "sections": [
            {
                "heading": "YTD Commission Summary",
                "content": "Year-to-date gross commission income of $956.5K across 3 closed deals totaling $36.3M. Current active pipeline of 5 deals totaling $144.2M with $3.5M in projected commission. Average commission rate of 2.8% across all transactions. On pace for $2M+ annual GCI — 65% ahead of last year's pace.",
                "key_metrics": [
                    {"label": "GCI YTD", "value": "$956.5K", "trend": "+65% YoY"},
                    {"label": "Closed Volume YTD", "value": "$36.3M", "trend": "3 deals"},
                    {"label": "Pipeline Value", "value": "$144.2M", "trend": "5 deals"},
                    {"label": "Projected Pipeline Comm.", "value": "$3.5M", "trend": "2.4% avg"},
                    {"label": "Projected Annual GCI", "value": "$2.0M+", "trend": "+65% YoY pace"},
                ],
            },
            {
                "heading": "Closed Transactions",
                "content": "Two closed sales in Q2: (1) Riverside Heights — $26.5M sale price, $662.5K commission, 2.5% rate. Closed May 1. (2) Poplar Pointe — $9.8M sale price, $294K commission, 3.0% rate. Closed May 15. Both transactions involved Linda Thornton (buyer side) and Thomas Garrett (seller side). Average time from LOI to close: 94 days.",
                "key_metrics": [
                    {"label": "Riverside Heights", "value": "$662.5K", "trend": "$26.5M, 2.5%"},
                    {"label": "Poplar Pointe", "value": "$294K", "trend": "$9.8M, 3.0%"},
                    {"label": "Avg LOI-to-Close", "value": "94 days", "trend": "Both Q2 deals"},
                    {"label": "Avg Commission", "value": "2.75%", "trend": "Blended rate"},
                ],
            },
            {
                "heading": "Active Pipeline Details",
                "content": "Village at Germantown ($24.5M, closing) — $735K commission, closing Q3. The Vue at Madison ($34M, due diligence) — $850K commission. Oakwood Crossings ($11.5M, LOI) — $345K commission. The Emerson ($68M, LOI) — $1.36M commission. Highland Ridge ($6.2M, underwriting) — $217K commission. Total: $144.2M, $3.5M commission.",
                "key_metrics": [
                    {"label": "Deal Count", "value": "5 active", "trend": "$144.2M total"},
                    {"label": "Largest Deal", "value": "$68M", "trend": "The Emerson — $1.36M commission"},
                    {"label": "Avg Active Deal", "value": "$28.8M", "trend": "Ranging $6.2-68M"},
                    {"label": "Commission at Risk", "value": "$3.5M", "trend": "Blended 2.4%"},
                ],
            },
            {
                "heading": "Business Development Activity",
                "content": "14 activities logged YTD: 6 property tours, 3 inspections/appraisals, 3 meetings/calls, 2 appraisals. 8 new listings taken totaling 1.14M SF. 4 repeat clients representing 40% of YTD revenue. Exclusive listings (3 of 8) commanding 3.5% average commission vs 2.3% for co-listings.",
                "key_metrics": [
                    {"label": "Activities YTD", "value": "14", "trend": "Tours, inspections, calls"},
                    {"label": "New Listings", "value": "8", "trend": "1.14M SF total"},
                    {"label": "Repeat Clients", "value": "4", "trend": "40% of YTD rev"},
                    {"label": "Exclusive Listings", "value": "3", "trend": "3.5% avg vs 2.3% co-list"},
                ],
            },
        ],
        "action_items": [
            "Prepare Q2 pipeline review presentation for managing director",
            "Activate backup offer strategy on The Emerson",
            "Follow up on Warburg Realty Trust for next acquisition mandate",
            "Schedule Q3 goal setting with Linda Thornton",
        ],
        "tags": ["commission", "report", "quarterly", "pipeline"],
    },
]


def _score(doc: dict[str, Any], q: str) -> int:
    """Simple keyword relevance score."""
    text = " ".join([
        doc.get("title", ""),
        doc.get("type", ""),
        " ".join(doc.get("tags", [])),
        " ".join(s.get("heading", "") for s in doc.get("sections", [])),
    ]).lower()
    words = q.lower().split()
    return sum(3 if w in doc.get("title", "").lower() else 1 for w in words if w in text)


class QueryRequest(BaseModel):
    query: str
    user_id: str | None = None


@app.get("/")
def root() -> dict[str, str]:
    return {"service": "mcp-document-generation", "status": "ok", "version": "0.3.0"}


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "healthy"}


@app.get("/mcp")
def mcp_info() -> dict[str, str]:
    return {"transport": "streamable-http", "service": "document_generation"}


@app.post("/mcp/query")
def mcp_query(payload: QueryRequest) -> dict[str, Any]:
    q = payload.query.lower()

    type_filter = None
    if any(w in q for w in ("portfolio summary", "portfolio performance", "portfolio overview")):
        type_filter = "portfolio_summary"
    elif any(w in q for w in ("offering memorandum", "om", "offering memo", "offering")):
        type_filter = "offering_memorandum"
    elif any(w in q for w in ("market survey", "market report", "market analysis", "market research", "comps")):
        type_filter = "market_survey"
    elif any(w in q for w in ("broker price opinion", "bpo", "price opinion", "valuation")):
        type_filter = "broker_price_opinion"
    elif any(w in q for w in ("commission report", "commission summary", "gci", "commission")):
        type_filter = "commission_report"

    candidates = [d for d in _DOCUMENT_TEMPLATES if (type_filter is None or d["type"] == type_filter)]

    scored = [(d, _score(d, q)) for d in candidates]
    scored.sort(key=lambda x: x[1], reverse=True)

    results = [d for d, s in scored if s > 0][:2]
    if not results:
        results = [d for d, s in scored[:1]]

    return {
        "service": "document_generation",
        "query": payload.query,
        "user_id": payload.user_id or "anonymous",
        "summary": f"Found {len(results)} document(s)",
        "documents": results,
    }


@app.get("/admin/data")
def admin_get_data() -> dict[str, Any]:
    return {
        "service": "document_generation",
        "total_templates": len(_DOCUMENT_TEMPLATES),
        "templates": [{"id": d["id"], "type": d["type"], "title": d["title"]} for d in _DOCUMENT_TEMPLATES],
    }


@app.post("/admin/data")
def admin_post_data(payload: dict[str, Any]) -> dict[str, Any]:
    doc = payload.get("document")
    if not doc or not isinstance(doc, dict):
        return {"error": "Missing 'document' field"}
    if "id" not in doc or "title" not in doc:
        return {"error": "Document must have 'id' and 'title'"}
    existing = [d for d in _DOCUMENT_TEMPLATES if d.get("id") == doc["id"]]
    if existing:
        _DOCUMENT_TEMPLATES[_DOCUMENT_TEMPLATES.index(existing[0])] = doc
        return {"service": "document_generation", "action": "updated", "id": doc["id"], "total": len(_DOCUMENT_TEMPLATES)}
    _DOCUMENT_TEMPLATES.append(doc)
    return {"service": "document_generation", "action": "added", "id": doc["id"], "total": len(_DOCUMENT_TEMPLATES)}


# ═══════════════════════════════════════════════════════════════════════
#  EXPORT — Models
# ═══════════════════════════════════════════════════════════════════════

class ExportRequest(BaseModel):
    type: str  # "report" | "table"
    format: str  # "xlsx" | "docx" | "pdf"
    title: str
    data: dict[str, Any]


def _report_sections(payload: ExportRequest) -> list[dict[str, Any]]:
    sections = payload.data.get("sections", [])
    if isinstance(sections, list) and sections:
        return [section for section in sections if isinstance(section, dict)]

    title = payload.title.strip() or "Generated document"
    return [
        {
            "heading": title,
            "content": "No structured sections were returned with this document.",
            "key_metrics": [],
        }
    ]


# ═══════════════════════════════════════════════════════════════════════
#  EXPORT — Generators
# ═══════════════════════════════════════════════════════════════════════

def _sanitize_name(title: str) -> str:
    name = re.sub(r"[^a-zA-Z0-9]+", "-", title).strip("-").lower()
    return name[:60] or "export"


def _generate_excel(payload: ExportRequest) -> bytes:
    import datetime

    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    generated = datetime.datetime.now(datetime.UTC).strftime("%B %d, %Y")

    wb = Workbook()
    ws = wb.active
    ws.title = payload.title[:31] or "Report"

    # ── Palette ────────────────────────────────────────────────────
    NAVY, SAPPHIRE, LIGHT, ALT, WHITE = "1F4E79", "2E75B6", "EBF3FB", "F7FAFD", "FFFFFF"

    # ── Reusable styles ────────────────────────────────────────────
    def _font(bold=False, size=9, color="2C2C2C", italic=False):
        return Font(bold=bold, size=size, color=color, name="Calibri", italic=italic)

    def _fill(color):
        return PatternFill(start_color=color, end_color=color, fill_type="solid")

    def _border(color="D0E4F5"):
        s = Side(style="thin", color=color)
        return Border(left=s, right=s, top=s, bottom=s)

    thin = _border()
    left_accent = Border(left=Side(style="medium", color=NAVY))
    center = Alignment(horizontal="center", vertical="center")
    vcenter = Alignment(vertical="center", indent=1)
    wrap = Alignment(wrap_text=True, vertical="top", indent=1)

    # ── Title header (rows 1–2) ─────────────────────────────────────
    for r, (val, fnt, ht) in enumerate([
        (payload.title,      _font(bold=True, size=15, color=WHITE),       36),
        (f"Generated {generated}  |  CONFIDENTIAL",
                              _font(size=8, color="8FBFE0", italic=True),  16),
    ], start=1):
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=5)
        c = ws.cell(row=r, column=1, value=val)
        c.font = fnt
        c.fill = _fill(NAVY)
        c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[r].height = ht

    row = 2

    if payload.type == "report":
        sections = _report_sections(payload)
        action_items = payload.data.get("action_items", [])
        tags = payload.data.get("tags", [])

        for sec in sections:
            # Section heading
            row += 2
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
            c = ws.cell(row=row, column=1, value=sec.get("heading", ""))
            c.font = _font(bold=True, size=11, color=NAVY)
            c.fill = _fill(LIGHT)
            c.border = left_accent
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[row].height = 22

            # Narrative content
            row += 1
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
            c = ws.cell(row=row, column=1, value=sec.get("content", ""))
            c.font = _font(size=9, color="444444")
            c.alignment = wrap
            ws.row_dimensions[row].height = 52

            # Key-metrics table
            metrics = sec.get("key_metrics", [])
            if metrics:
                row += 1
                for col, label in enumerate(["Metric", "Value", "Trend"], 1):
                    c = ws.cell(row=row, column=col, value=label)
                    c.font = _font(bold=True, size=9, color=WHITE)
                    c.fill = _fill(SAPPHIRE)
                    c.border = thin
                    c.alignment = center
                ws.row_dimensions[row].height = 18

                for i, m in enumerate(metrics):
                    row += 1
                    row_fill = _fill(ALT) if i % 2 == 0 else _fill(WHITE)
                    for col, (val, bold) in enumerate(
                        [(m.get("label",""), False), (m.get("value",""), True), (m.get("trend",""), False)], 1
                    ):
                        c = ws.cell(row=row, column=col, value=val)
                        c.font = _font(bold=bold, size=9)
                        c.fill = row_fill
                        c.border = thin
                        c.alignment = vcenter
                    ws.row_dimensions[row].height = 16

        if action_items:
            row += 2
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
            c = ws.cell(row=row, column=1, value="Action Items")
            c.font = _font(bold=True, size=11, color=NAVY)
            c.fill = _fill(LIGHT)
            c.border = left_accent
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[row].height = 22

            for i, item in enumerate(action_items):
                row += 1
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
                c = ws.cell(row=row, column=1, value=f"  ▸  {item}")
                c.font = _font(size=9, color="333333")
                c.fill = _fill(ALT) if i % 2 == 0 else _fill(WHITE)
                c.alignment = Alignment(wrap_text=True, vertical="top")
                ws.row_dimensions[row].height = 16

        if tags:
            row += 2
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
            c = ws.cell(row=row, column=1, value="Tags")
            c.font = _font(bold=True, size=11, color=NAVY)
            c.fill = _fill(LIGHT)
            c.border = left_accent
            c.alignment = Alignment(horizontal="left", vertical="center", indent=1)
            ws.row_dimensions[row].height = 22
            row += 1
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=5)
            c = ws.cell(row=row, column=1, value="  " + "  •  ".join(tags))
            c.font = _font(size=9, color="555555")
            ws.row_dimensions[row].height = 16

        ws.column_dimensions["A"].width = 32
        ws.column_dimensions["B"].width = 20
        ws.column_dimensions["C"].width = 28
        ws.column_dimensions["D"].width = 10
        ws.column_dimensions["E"].width = 10

    else:
        headers = payload.data.get("headers", [])
        rows_data = payload.data.get("rows", [])

        if headers:
            row += 2
            for col, h in enumerate(headers, 1):
                c = ws.cell(row=row, column=col, value=h)
                c.font = _font(bold=True, size=9, color=WHITE)
                c.fill = _fill(SAPPHIRE)
                c.border = thin
                c.alignment = center
                ws.column_dimensions[chr(64 + col) if col <= 26 else "A"].width = 22
            ws.row_dimensions[row].height = 20

        for i, row_data in enumerate(rows_data):
            row += 1
            row_fill = _fill(ALT) if i % 2 == 0 else _fill(WHITE)
            for col, val in enumerate(row_data, 1):
                c = ws.cell(row=row, column=col, value=str(val))
                c.font = _font(size=9)
                c.fill = row_fill
                c.border = thin
                c.alignment = vcenter
            ws.row_dimensions[row].height = 16

    # Freeze below the 2-row branded header
    ws.freeze_panes = ws.cell(row=3, column=1)
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generate_docx(payload: ExportRequest) -> bytes:
    import datetime

    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    NAVY = RGBColor(0x1F, 0x4E, 0x79)
    SAPPHIRE = RGBColor(0x2E, 0x75, 0xB6)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    GREY = RGBColor(0x99, 0x99, 0x99)

    generated = datetime.datetime.now(datetime.UTC).strftime("%B %d, %Y")

    doc = Document()

    # ── Margins ─────────────────────────────────────────────────────
    pg = doc.sections[0]
    pg.left_margin = pg.right_margin = Cm(2.54)
    pg.top_margin = pg.bottom_margin = Cm(2.54)

    # ── Base styles ─────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    for lvl, size, space_before in [(1, 22, 0), (2, 13, 14)]:
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = NAVY
        s.paragraph_format.space_before = Pt(space_before)
        s.paragraph_format.space_after = Pt(4)

    # ── XML helpers ─────────────────────────────────────────────────
    def _shade_cell(cell: Any, fill_hex: str) -> None:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    def _left_border(para: Any, color: str = "1F4E79", sz: int = 24) -> None:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(sz))
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), color)
        pBdr.append(left)
        pPr.append(pBdr)

    def _page_number_field(run: Any) -> None:
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)

    # ── Header ──────────────────────────────────────────────────────
    hdr_section = doc.sections[0].header
    for p in hdr_section.paragraphs:
        p.clear()
    hdr_para = hdr_section.paragraphs[0] if hdr_section.paragraphs else hdr_section.add_paragraph()
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hdr_para.add_run(payload.title)
    hr.font.size = Pt(8)
    hr.font.color.rgb = GREY
    hr.font.italic = True
    pPr = hdr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "CCCCCC")
    pBdr.append(bot)
    pPr.append(pBdr)

    # ── Footer with page number ──────────────────────────────────────
    ftr_section = doc.sections[0].footer
    for p in ftr_section.paragraphs:
        p.clear()
    ftr_para = ftr_section.paragraphs[0] if ftr_section.paragraphs else ftr_section.add_paragraph()
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    txt_run = ftr_para.add_run(f"{generated}  •  CONFIDENTIAL  •  Page ")
    txt_run.font.size = Pt(8)
    txt_run.font.color.rgb = GREY
    pn_run = ftr_para.add_run()
    pn_run.font.size = Pt(8)
    pn_run.font.color.rgb = GREY
    _page_number_field(pn_run)

    # ── Title block ─────────────────────────────────────────────────
    title_para = doc.add_heading(payload.title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_para.add_run(f"{generated}  •  CONFIDENTIAL")
    dr.font.size = Pt(10)
    dr.font.color.rgb = SAPPHIRE
    date_para.paragraph_format.space_after = Pt(20)

    # ── Helper: styled metric/data table ────────────────────────────
    def _build_table(rows_data: list[list[str]], headers: list[str]) -> Any:
        ncols = max(len(headers), max((len(r) for r in rows_data), default=0), 1)
        tbl = doc.add_table(rows=1 + len(rows_data), cols=ncols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            _shade_cell(hdr_cells[i], "1F4E79")
            for p in hdr_cells[i].paragraphs:
                run = p.add_run(h)
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
                p.paragraph_format.space_after = Pt(0)

        for r_i, row_vals in enumerate(rows_data):
            fill = "EBF3FB" if r_i % 2 == 0 else "FFFFFF"
            data_cells = tbl.rows[r_i + 1].cells
            for ci, val in enumerate(row_vals):
                if ci < len(data_cells):
                    _shade_cell(data_cells[ci], fill)
                    for p in data_cells[ci].paragraphs:
                        run = p.add_run(val)
                        run.font.size = Pt(9)
                        p.paragraph_format.space_after = Pt(0)
        return tbl

    # ── Content ─────────────────────────────────────────────────────
    if payload.type == "report":
        sections = _report_sections(payload)
        action_items = payload.data.get("action_items", [])
        tags = payload.data.get("tags", [])

        for sec in sections:
            h2 = doc.add_heading(sec.get("heading", ""), level=2)
            _left_border(h2)
            doc.add_paragraph(sec.get("content", ""))

            metrics = sec.get("key_metrics", [])
            if metrics:
                _build_table(
                    [[m.get("label",""), m.get("value",""), m.get("trend","")] for m in metrics],
                    ["Metric", "Value", "Trend"],
                )
                doc.add_paragraph().paragraph_format.space_after = Pt(4)

        if action_items:
            h2 = doc.add_heading("Action Items", level=2)
            _left_border(h2)
            for item in action_items:
                doc.add_paragraph(item, style="List Bullet")

        if tags:
            h2 = doc.add_heading("Tags", level=2)
            _left_border(h2)
            doc.add_paragraph("  ".join(f"[{t}]" for t in tags))

    else:
        headers = payload.data.get("headers", [])
        rows_data = payload.data.get("rows", [])
        if headers:
            _build_table([[str(v) for v in row] for row in rows_data], headers)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generate_pdf(payload: ExportRequest) -> bytes:
    import datetime

    from weasyprint import HTML

    generated = datetime.datetime.now(datetime.UTC).strftime("%B %d, %Y")

    if payload.type == "report":
        sections = _report_sections(payload)
        action_items = payload.data.get("action_items", [])
        table_headers: list = []
        table_rows: list = []
    else:
        sections = []
        action_items = []
        table_headers = payload.data.get("headers", [])
        table_rows = payload.data.get("rows", [])

    def _kpi_grid(metrics: list) -> str:
        if not metrics:
            return ""
        cols = 3
        html = '<table style="width:100%;border-collapse:separate;border-spacing:5px;margin:10px 0 0 0;">'
        for i in range(0, len(metrics), cols):
            batch = metrics[i : i + cols]
            html += "<tr>"
            for m in batch:
                label = m.get("label", "")
                value = m.get("value", "")
                trend = m.get("trend", "")
                t = trend.lstrip()
                t_color = "#2E7D32" if t.startswith("+") or "↑" in t else (
                    "#C62828" if t.startswith("-") or "↓" in t else "#666"
                )
                html += (
                    f'<td style="width:33%;border:1px solid #D0E4F5;border-top:3px solid #2E75B6;'
                    f'padding:8px 12px;background:#F7FAFD;vertical-align:top;">'
                    f'<div style="font-size:7pt;color:#888;text-transform:uppercase;'
                    f'letter-spacing:0.6pt;margin-bottom:3px;">{label}</div>'
                    f'<div style="font-size:14pt;font-weight:bold;color:#1F4E79;'
                    f'margin-bottom:2px;">{value}</div>'
                    f'<div style="font-size:8pt;color:{t_color};">{trend}</div>'
                    f'</td>'
                )
            for _ in range(cols - len(batch)):
                html += "<td></td>"
            html += "</tr>"
        html += "</table>"
        return html

    # Sections
    sections_html = ""
    for s in sections:
        heading = s.get("heading", "")
        content = s.get("content", "")
        sections_html += (
            f'<div style="page-break-inside:avoid;margin-bottom:22px;">'
            f'<div style="border-left:4px solid #1F4E79;padding:2px 0 2px 10px;margin-bottom:6px;">'
            f'<h2 style="color:#1F4E79;font-size:12pt;margin:0;font-weight:bold;">{heading}</h2>'
            f'</div>'
            f'<p style="font-size:9.5pt;line-height:1.65;color:#333;margin:0 0 4px 0;">{content}</p>'
            f'{_kpi_grid(s.get("key_metrics", []))}'
            f'</div>'
        )

    # Data table
    table_html = ""
    if table_headers:
        th = "".join(
            f'<th style="padding:6px 10px;border:1px solid #1a4068;background:#1F4E79;'
            f'color:#fff;font-size:9pt;text-align:left;">{h}</th>'
            for h in table_headers
        )
        trs = ""
        for i, row in enumerate(table_rows):
            bg = "#F7FAFD" if i % 2 else "#FFFFFF"
            td = "".join(
                f'<td style="padding:5px 10px;border:1px solid #D5E6F3;font-size:9pt;background:{bg};">{str(v)}</td>'
                for v in row
            )
            trs += f"<tr>{td}</tr>"
        table_html = (
            f'<table style="width:100%;border-collapse:collapse;margin-top:10px;">'
            f'<thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>'
        )

    # Action items
    actions_html = ""
    if action_items:
        items = "".join(
            f'<li style="font-size:9.5pt;margin-bottom:5px;line-height:1.5;">{item}</li>'
            for item in action_items
        )
        actions_html = (
            f'<div style="page-break-inside:avoid;margin-bottom:22px;">'
            f'<div style="border-left:4px solid #1F4E79;padding:2px 0 2px 10px;margin-bottom:6px;">'
            f'<h2 style="color:#1F4E79;font-size:12pt;margin:0;font-weight:bold;">Action Items</h2>'
            f'</div>'
            f'<ul style="margin:6px 0 0 0;padding-left:20px;">{items}</ul>'
            f'</div>'
        )

    html_str = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
@page {{
    margin: 2.2cm 2cm 2.5cm 2cm;
    @bottom-right {{
        content: "Page " counter(page) " of " counter(pages);
        font-family: 'Liberation Sans', Arial, sans-serif;
        font-size: 7.5pt;
        color: #aaa;
    }}
    @bottom-center {{
        content: "CONFIDENTIAL";
        font-family: 'Liberation Sans', Arial, sans-serif;
        font-size: 7pt;
        color: #aaa;
        letter-spacing: 1.5pt;
    }}
}}
body {{
    font-family: 'Liberation Sans', 'Arial', sans-serif;
    color: #2C2C2C;
    margin: 0;
    font-size: 10pt;
}}
</style>
</head>
<body>
<div style="background:#1F4E79;color:#fff;padding:22px 26px 18px 26px;margin-bottom:26px;">
    <div style="font-size:7pt;letter-spacing:2pt;text-transform:uppercase;color:#8FBFE0;margin-bottom:7px;">CONFIDENTIAL REPORT</div>
    <div style="font-size:19pt;font-weight:bold;line-height:1.2;margin-bottom:8px;">{payload.title}</div>
    <div style="font-size:8.5pt;color:#8FBFE0;">Generated {generated}</div>
</div>
{sections_html}
{table_html}
{actions_html}
</body>
</html>"""

    pdf_bytes = HTML(string=html_str).write_pdf()
    if not pdf_bytes:
        raise RuntimeError(
            "weasyprint produced an empty PDF. "
            "Check that fontconfig and system fonts are installed in the container."
        )
    return pdf_bytes


# ═══════════════════════════════════════════════════════════════════════
#  EXPORT — Endpoints
# ═══════════════════════════════════════════════════════════════════════

FORMAT_MAP = {
    "xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx"),
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx"),
    "pdf": ("application/pdf", ".pdf"),
}


@app.post("/export")
def export_document(payload: ExportRequest) -> Response:
    fmt_info = FORMAT_MAP.get(payload.format)
    if not fmt_info:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {payload.format}")
    content_type, ext = fmt_info

    try:
        if payload.format == "xlsx":
            content = _generate_excel(payload)
        elif payload.format == "docx":
            content = _generate_docx(payload)
        elif payload.format == "pdf":
            content = _generate_pdf(payload)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {payload.format}")
    except ImportError as e:
        raise HTTPException(
            status_code=500,
            detail=f"Export library not available: {e.name}. Run: pip install {e.name}",
        ) from e
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export generation failed: {e}") from e

    filename = f"{_sanitize_name(payload.title)}{ext}"
    return Response(
        content=content,
        media_type=content_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
